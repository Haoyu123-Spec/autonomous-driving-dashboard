#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成训练分析报告 Word 文档"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

style = doc.styles["Normal"]
style.font.name = "微软雅黑"
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "微软雅黑"
    hs.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

title = doc.add_heading("Dueling DQN 多AGV碰撞避免 — 训练分析报告", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    "模型：Dueling DQN | 场景：4 AGV, 10x10 网格世界 | "
    "训练：3000 集 | 动作空间：{Stay, Right, Left, Up, Down}",
).alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("")

# ═══════════════════════════
# 图1：训练曲线
# ═══════════════════════════
doc.add_heading("图1：训练曲线", level=1)

doc.add_paragraph(
    "四合一训练曲线图，展示奖励、损失、步数、探索率随训练集数的变化趋势。"
)

table1 = doc.add_table(rows=5, cols=3, style="Light Grid Accent 1")
table1.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, h in enumerate(["子图", "内容", "解读"]):
    cell = table1.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(10)

data1 = [
    ["左上：奖励曲线",
     "单集得分与最近100集滑动平均",
     "avg100 从初始负值持续上升，最终稳定在 +20~+24 区间，说明策略持续改善。avg100 越过0线意味着平均奖励由负转正，策略开始生效。"],
    ["右上：损失曲线",
     "SmoothL1Loss",
     "损失稳定在 0.1~0.5 之间，无发散或持续上升趋势。训练后期损失不再明显下降说明模型已收敛到当前架构的能力上限。SmoothL1Loss 有效避免了 Q 值爆炸问题。"],
    ["左下：每集步数",
     "每集完成的步数",
     "初始约200步，策略改善后降至50~150步。步数减少意味着 AGV 更高效地到达目标，而非因碰撞提前终止。"],
    ["右下：探索率衰减",
     "epsilon 随 Episode 变化",
     "从1.0按 episode 衰减至0.05，约在1500集后接近底值。衰减曲线平滑，证明 per-episode 衰减已正确生效。epsilon_min=0.05 保留了最小探索能力。"],
]

for r, row in enumerate(data1, 1):
    for c, text in enumerate(row):
        cell = table1.rows[r].cells[c]
        cell.text = text
        for p in cell.paragraphs:
            p.runs[0].font.size = Pt(10)

doc.add_paragraph("")

p = doc.add_paragraph()
run = p.add_run("[总评] ")
run.bold = True
run.font.size = Pt(10)
run = p.add_run(
    "训练健康，策略在持续学习改善，没有发散或过拟合迹象。"
    "四大核心修复（per-episode epsilon衰减、SmoothL1Loss、奖励缩放、梯度裁剪+硬更新）全部生效。"
)
run.font.size = Pt(10)

doc.add_paragraph("")

# ═══════════════════════════
# 图2：Q值分析
# ═══════════════════════════
doc.add_heading("图2：Q值分布与价值流分解", level=1)

doc.add_paragraph(
    "Dueling DQN 将 Q(s,a) 分解为状态价值 V(s) 与优势函数 A(s,a)。"
    "本图通过三个直方图展示模型学到的 Q 值、V 值和 Advantage 分布。"
)

table2 = doc.add_table(rows=4, cols=2, style="Light Grid Accent 1")
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, h in enumerate(["子图", "解读"]):
    cell = table2.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(10)

data2 = [
    ["左：各动作 Q 值分布",
     "五个动作的 Q 值分布高度重叠，均值都在 ~2.53 附近。这是 Dueling DQN 的预期行为：V(s) 主导 Q 值，A(s,a) 接近 0。但分布几乎完全相同意味着策略对不同动作的偏好很弱，对复杂场景的决策区分能力不足。"],
    ["中：状态价值 V(s) 分布",
     "V(s) 呈近似正态分布，集中在 2.5 左右。说明网络学到了有意义的状态估值，但范围较窄（约 2.2~2.8），可能是因为 10x10 小世界在 4 AGV 场景下状态多样性有限。"],
    ["右：优势函数 A(s,a) 分布",
     "所有动作的 A(s,a) 集中在 0 附近（+-0.15 以内），验证了 Dueling 架构的 A_mean=0 约束。Advantage 区分度很小，印证策略过早收敛到安全但不一定最优的行为模式（频繁 Stay）。"],
]

for r, row in enumerate(data2, 1):
    for c, text in enumerate(row):
        cell = table2.rows[r].cells[c]
        cell.text = text
        for p in cell.paragraphs:
            p.runs[0].font.size = Pt(10)

doc.add_paragraph("")

p = doc.add_paragraph()
run = p.add_run("[总评] ")
run.bold = True
run.font.size = Pt(10)
run = p.add_run(
    "Dueling 架构的 V/A 分解正常工作，符合理论预期。"
    "但 Q 值区分度不足是当前模型的核心局限——五个动作的 Q 值几乎相同，"
    "说明策略缺乏对动作后果的精细判断。"
    "下一步优化方向：引入更丰富奖励信号（时间惩罚、路径效率奖励）、"
    "延长探索期、或采用 PER（优先经验回放）让网络接触更多高学习价值的稀有样本。"
)
run.font.size = Pt(10)

doc.add_paragraph("")

# ═══════════════════════════
# 图3：轨迹可视化
# ═══════════════════════════
doc.add_heading("图3：策略轨迹可视化", level=1)

doc.add_paragraph(
    "左图为 AGV 在一个完整 Episode 中的路径轨迹，"
    "右图为该集中各动作的使用次数分布。"
)

table3 = doc.add_table(rows=5, cols=2, style="Light Grid Accent 1")
table3.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, h in enumerate(["子图", "解读"]):
    cell = table3.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(10)

data3 = [
    ["左：AGV 路径轨迹",
     "3/4 个 AGV 成功到达目标：AGV0（蓝）、AGV1（橙）、AGV3（红）均已抵达；AGV2（绿）未能到达。路径无交叉碰撞，说明碰撞避免策略有效。轨迹平滑，AGV 学会了先移开再朝目标走的基本策略。"],
    ["右：动作分布",
     "Stay（停留）占比最高（约500次），其次为 Right 和 Up。Stay 主导反映出：① 4个AGV中有1个未到达，可能被锁死；② 策略过度保守，即使在安全区域也不愿移动。这印证了 Q 值分析中动作同质化的结论。"],
    ["底部：奖励信息",
     "3 个 AGV 获得正奖励（~14~15），AGV2 获得负奖励——因没到达目标、步数耗尽产生累积负奖励。碰撞数=0 说明本集无碰撞发生，避碰策略有效。"],
    ["关键发现",
     "Stay 占比过高和 Q 值区分度不足是同一问题的两面。优化方向：① 加大 reward_shaping_scale 让距离变化奖励更强；② 引入 idle_penalty 鼓励移动；③ 增加训练场景多样性迫使网络学习更精细的动作区分。"],
]

for r, row in enumerate(data3, 1):
    for c, text in enumerate(row):
        cell = table3.rows[r].cells[c]
        cell.text = text
        for p in cell.paragraphs:
            p.runs[0].font.size = Pt(10)

doc.add_paragraph("")

p = doc.add_paragraph()
run = p.add_run("[总评] ")
run.bold = True
run.font.size = Pt(10)
run = p.add_run(
    "避碰策略有效，3/4 到达率可接受。Stay 占比过高和 Q 值区分度不足指向同一个根因——"
    "需要更强的奖励信号来区分积极行动与原地等待。下一步优化应聚焦于鼓励更积极的移动策略。"
)
run.font.size = Pt(10)

doc.add_paragraph("")

# ═══════════════════════════
# 总结
# ═══════════════════════════
doc.add_heading("总结与下一步建议", level=1)

items = [
    ("训练修复成功",
     "四大修复（per-episode epsilon衰减、SmoothL1Loss、奖励缩放、梯度裁剪+硬更新）全部生效，"
     "avg100 从修复前的 -87 提升到 +24，训练稳定无发散。"),
    ("Q 值区分度不足",
     "5 个动作的 Q 值分布几乎相同（均值 ~2.53），Advantage 集中在 +-0.15，"
     "说明策略对不同动作的偏好很弱。需要更强的奖励信号和更多探索来打破动作同质化。"),
    ("Stay 占比过高",
     "Stay 动作占所有动作的 50% 以上，策略过度保守。"
     "可引入 idle_penalty、增大 shaping reward、或使用 NoisyNet 替代 epsilon-greedy 来增加探索多样性。"),
    ("推荐下一步",
     "尝试训练 rl_advanced 版本：该版本包含 6 阶段课程学习、动态障碍物、三区协同奖励、"
     "方向引导奖励、紧急订单插入、Self-Attention 和 PER，复杂度更高，有望学习到更精细的决策策略。"),
]

for title_text, desc in items:
    p = doc.add_paragraph()
    run = p.add_run(f"[{title_text}] ")
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(desc)
    run.font.size = Pt(11)

save_path = "e:/carAI/训练分析报告_Dueling_DQN.docx"
doc.save(save_path)
print(f"报告已保存: {save_path}")
